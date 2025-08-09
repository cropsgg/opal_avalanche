from __future__ import annotations

from typing import List, Dict, Any
import io
import re
from docx import Document


def extract_paragraphs(docx_bytes: bytes) -> List[str]:
    """Legacy function - extract paragraph text only"""
    f = io.BytesIO(docx_bytes)
    d = Document(f)
    return [p.text for p in d.paragraphs]


def extract_text_with_paras(docx_bytes: bytes) -> List[Dict[str, Any]]:
    """Extract paragraphs with metadata and formatting information"""
    f = io.BytesIO(docx_bytes)
    d = Document(f)
    
    paragraphs = []
    para_id = 1
    
    for p in d.paragraphs:
        text = p.text.strip()
        
        if not text or len(text) < 10:  # Skip empty or very short paragraphs
            continue
            
        # Analyze paragraph style and formatting
        style_name = p.style.name if p.style else "Normal"
        
        # Check if it's a heading
        is_heading = (
            "heading" in style_name.lower() or
            p.style.name.startswith("Heading") if p.style else False
        )
        
        # Extract heading level
        heading_level = None
        if is_heading and p.style:
            level_match = re.search(r'(\d+)', p.style.name)
            if level_match:
                heading_level = int(level_match.group(1))
        
        # Check for numbered paragraphs
        numbered_match = re.match(r'^(\d+\.?\s*)', text)
        
        # Analyze text formatting for additional metadata
        runs_info = []
        for run in p.runs:
            if run.text.strip():
                runs_info.append({
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline
                })
        
        paragraphs.append({
            "para_id": para_id,
            "text": text,
            "style": style_name,
            "is_heading": is_heading,
            "heading_level": heading_level,
            "is_numbered": bool(numbered_match),
            "number": numbered_match.group(1).strip() if numbered_match else None,
            "word_count": len(text.split()),
            "char_count": len(text),
            "runs": runs_info
        })
        
        para_id += 1
    
    return paragraphs


def extract_headings(docx_bytes: bytes) -> List[Dict[str, Any]]:
    """Extract document headings with hierarchy"""
    paragraphs = extract_text_with_paras(docx_bytes)
    headings = []
    
    for para in paragraphs:
        if para["is_heading"]:
            headings.append({
                "para_id": para["para_id"],
                "text": para["text"],
                "level": para["heading_level"] or 1,
                "style": para["style"]
            })
    
    return headings


def extract_tables(docx_bytes: bytes) -> List[Dict[str, Any]]:
    """Extract table content from DOCX"""
    f = io.BytesIO(docx_bytes)
    d = Document(f)
    
    tables = []
    table_id = 1
    
    for table in d.tables:
        table_data = []
        
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)
        
        if table_data:  # Only add non-empty tables
            tables.append({
                "table_id": table_id,
                "rows": len(table_data),
                "cols": len(table_data[0]) if table_data else 0,
                "data": table_data
            })
            table_id += 1
    
    return tables


